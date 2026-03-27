from llm_handlers.api_handler import get_llm_response
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
import io
import re

def generate_mind_map(params):
    """Generate lesson mind map document"""
    
    prompt = _build_mind_map_prompt(params)
    
    try:
        # Get content from LLM
        content = get_llm_response(prompt, params["llm_config"])
        
        # Create Word document
        docx_file = _create_mind_map_docx(content, params)
        
        return {
            "success": True,
            "content": content,
            "docx_file": docx_file
        }
        
    except Exception as e:
        return {"success": False, "error": str(e)}

def _build_mind_map_prompt(params):
    """Build prompt for mind map generation"""
    
    main_branches = params.get("main_branches") or 6
    depth_levels = params.get("depth_levels") or 3
    include_examples = params.get("include_examples", True)
    highlight_hierarchy = params.get("highlight_hierarchy", True)
    hierarchy_instruction = "Emphasize hierarchical relationships between concepts."
    if not highlight_hierarchy:
        hierarchy_instruction = "Keep the hierarchy minimal and focus on main branches."
    
    prompt = f"""
    Create a lesson mind map for {params['subject']} at {params['grade_level']} level.
    
    Topic: {params['topic']}
    Main branches: {main_branches}
    Depth levels: {depth_levels}
    Include examples/applications: {include_examples}
    Hierarchy emphasis: {highlight_hierarchy}
    
    FORMATTING RULES:
    - Use Markdown only
    - Start with a single H1 title (#) for the central topic
    - Use H2 headings (##) for each main branch
    - Use nested bullet points for sub-branches with 2 spaces per level
    - Keep each node short (max 8-10 words)
    - Avoid paragraphs and horizontal rules
    - Write in the same language as the subject/topic
    - {hierarchy_instruction}
    """
    
    return prompt

def _create_mind_map_docx(content, params):
    """Create Word document from mind map content"""
    
    doc = Document()
    
    title_text = f"{params['subject']} - {params.get('doc_type', 'Lesson Mind Map')}"
    title = doc.add_heading(title_text, 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_heading(f"Topic: {params['topic']}", level=2)
    doc.add_heading(f"Grade Level: {params['grade_level']}", level=3)
    
    _add_mind_map_content_to_docx(doc, content)
    
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io.getvalue()

def _add_mind_map_content_to_docx(doc, content):
    """Add mind map content to Word document, preserving hierarchy"""
    lines = content.split('\n')
    
    for line in lines:
        raw_line = line.rstrip('\n')
        stripped = raw_line.strip()
        
        if not stripped:
            doc.add_paragraph("")
            continue
        
        if stripped.startswith('#'):
            hash_count = 0
            for char in stripped:
                if char == '#':
                    hash_count += 1
                else:
                    break
            
            heading_text = stripped[hash_count:].strip()
            heading_text = re.sub(r'\*\*(.*?)\*\*', r'\1', heading_text)
            
            actual_level = min(hash_count, 4)
            if actual_level > 0 and heading_text:
                doc.add_heading(heading_text, level=actual_level)
                continue
        
        bullet_line = raw_line.lstrip()
        if bullet_line.startswith('- ') or bullet_line.startswith('* ') or bullet_line.startswith('• '):
            indent_spaces = len(raw_line) - len(bullet_line)
            level = max(0, indent_spaces // 2)
            text = bullet_line[2:].strip()
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            paragraph = doc.add_paragraph(text, style='List Bullet')
            if level > 0:
                paragraph.paragraph_format.left_indent = Inches(0.25 * level)
            continue
        
        if len(stripped) > 2 and stripped[0].isdigit() and stripped[1:3] in ['. ', ') ']:
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', stripped[3:].strip())
            doc.add_paragraph(text, style='List Number')
        else:
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', stripped)
            doc.add_paragraph(text)
