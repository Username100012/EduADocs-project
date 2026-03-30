from llm_handlers.api_handler import get_llm_response
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io
import re

def generate_lesson_plan(params):
    """Generate lesson plan document"""
    
    prompt = _build_lesson_plan_prompt(params)
    
    try:
        # Get content from LLM
        content = get_llm_response(prompt, params["llm_config"])
        
        # Create Word document
        docx_file = _create_lesson_plan_docx(content, params)
        
        return {
            "success": True,
            "content": content,
            "docx_file": docx_file
        }
        
    except Exception as e:
        return {"success": False, "error": str(e)}

def _build_lesson_plan_prompt(params):
    """Build prompt for lesson plan generation"""
    
    duration_minutes = params.get("duration_minutes")
    learning_objectives = params.get("learning_objectives") or "Not specified"
    materials = params.get("materials") or "Not specified"
    methodology = params.get("methodology") or "Not specified"
    assessment_strategy = params.get("assessment_strategy") or "Not specified"
    lesson_flow = params.get("lesson_flow") or "Not specified"
    include_differentiation = params.get("include_differentiation", False)
    
    prompt = f"""
    Create a detailed lesson plan for {params['subject']} at {params['grade_level']} level.
    
    Topic: {params['topic']}
    Duration: {duration_minutes} minutes
    Learning objectives: {learning_objectives}
    Materials/Resources: {materials}
    Teaching methodology: {methodology}
    Assessment strategy: {assessment_strategy}
    Lesson flow: {lesson_flow}
    Include differentiation/adaptations: {include_differentiation}
    
    Structure the lesson plan with:
    1. Lesson title and objectives
    2. Prior knowledge or prerequisites
    3. Step-by-step lesson flow (warm-up, main activity, closure), using the provided flow when specified
    4. Materials and resources list
    5. Assessment strategy
    6. Differentiation/adaptations for diverse learners (if requested)
    7. Homework or extension activities (optional)
    
    FORMATTING RULES:
    - Use Markdown headings (#, ##, ###) for sections
    - Use bullet points for lists and activities
    - Keep headings consistent and avoid horizontal rules
    - Write in the same language as the subject/topic
    """
    
    return prompt

def _create_lesson_plan_docx(content, params):
    """Create Word document from lesson plan content"""
    
    doc = Document()
    
    title_text = f"{params['subject']} - {params.get('doc_type', 'Lesson Plan')}"
    title = doc.add_heading(title_text, 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_heading(f"Topic: {params['topic']}", level=2)
    doc.add_heading(f"Grade Level: {params['grade_level']}", level=3)
    
    if params.get("duration_minutes"):
        doc.add_paragraph(f"Duration: {params['duration_minutes']} minutes")
    
    _add_formatted_content_to_docx(doc, content)
    
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io.getvalue()

def _add_formatted_content_to_docx(doc, content):
    """Add formatted content to Word document, parsing Markdown formatting"""
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph("")
            continue
        
        if line.startswith('#'):
            hash_count = 0
            for char in line:
                if char == '#':
                    hash_count += 1
                else:
                    break
            
            heading_text = line[hash_count:].strip()
            heading_text = re.sub(r'\*\*(.*?)\*\*', r'\1', heading_text)
            
            actual_level = min(hash_count, 4)
            if actual_level > 0 and heading_text:
                doc.add_heading(heading_text, level=actual_level)
                continue
        
        if line.startswith('- ') or line.startswith('* ') or line.startswith('• '):
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', line[2:].strip())
            doc.add_paragraph(text, style='List Bullet')
        elif len(line) > 2 and line[0].isdigit() and line[1:3] in ['. ', ') ']:
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', line[3:].strip())
            doc.add_paragraph(text, style='List Number')
        else:
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            doc.add_paragraph(text)
