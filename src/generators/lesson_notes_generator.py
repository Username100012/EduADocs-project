from llm_handlers.api_handler import get_llm_response
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io
import re

def generate_lecture_notes(params):
    """Generate lecture notes document"""
    
    prompt = _build_lecture_notes_prompt(params)
    
    try:
        # Get content from LLM
        content = get_llm_response(prompt, params["llm_config"])
        
        # Create Word document
        docx_file = _create_lecture_notes_docx(content, params)
        
        return {
            "success": True,
            "content": content,
            "docx_file": docx_file
        }
        
    except Exception as e:
        return {"success": False, "error": str(e)}

def _build_lecture_notes_prompt(params):
    """Build prompt for lecture notes generation"""
    
    detail_level = params.get("detail_level") or "Standard"
    format_style = params.get("format_style") or "Paragraphs"
    include_examples = params.get("include_examples", True)
    include_references = params.get("include_references", False)
    
    prompt = f"""
    Create lecture notes for {params['subject']} at {params['grade_level']} level.
    
    Topic: {params['topic']}
    Detail level: {detail_level}
    Format style: {format_style}
    Include examples: {include_examples}
    Include references: {include_references}
    
    Structure the notes with:
    1. Brief introduction and learning goals
    2. Key concepts and definitions
    3. Explanations with examples (if requested)
    4. Key takeaways or summary
    5. References or further reading (if requested)
    
    FORMATTING RULES:
    - Use Markdown headings (#, ##, ###) for sections
    - Use bullet points when the format style is "Bullet Points" or "Outline"
    - Keep the content clear and classroom-ready
    - Write in the same language as the subject/topic
    """
    
    return prompt

def _create_lecture_notes_docx(content, params):
    """Create Word document from lecture notes content"""
    
    doc = Document()
    
    title_text = f"{params['subject']} - {params.get('doc_type', 'Lecture Notes')}"
    title = doc.add_heading(title_text, 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_heading(f"Topic: {params['topic']}", level=2)
    doc.add_heading(f"Grade Level: {params['grade_level']}", level=3)
    
    _add_formatted_content_to_docx(doc, content)
    
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io.getvalue()

def _add_formatted_content_to_docx(doc, content):
    """Add formatted content to Word document, parsing Markdown formatting"""
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph("")
            continue
        
        if line.startswith('#'):
            hash_count = 0
            for char in line:
                if char == '#':
                    hash_count += 1
                else:
                    break
            
            heading_text = line[hash_count:].strip()
            heading_text = re.sub(r'\*\*(.*?)\*\*', r'\1', heading_text)
            
            actual_level = min(hash_count, 4)
            if actual_level > 0 and heading_text:
                doc.add_heading(heading_text, level=actual_level)
                continue
        
        if line.startswith('- ') or line.startswith('* ') or line.startswith('• '):
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', line[2:].strip())
            doc.add_paragraph(text, style='List Bullet')
        elif len(line) > 2 and line[0].isdigit() and line[1:3] in ['. ', ') ']:
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', line[3:].strip())
            doc.add_paragraph(text, style='List Number')
        else:
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            doc.add_paragraph(text)
